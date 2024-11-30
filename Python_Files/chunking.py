from transformers import GPT2TokenizerFast
import os
import re
from typing import List, Dict
import unicodedata
from datetime import datetime
import time
from docx import Document

class TextCleaner:
    @staticmethod
    def remove_extra_whitespace(text: str) -> str:
        """Remove extra whitespace, including multiple newlines."""
        # Replace multiple spaces with single space
        text = re.sub(r'\s+', ' ', text)
        # Replace multiple newlines with double newline
        text = re.sub(r'\n\s*\n\s*', '\n\n', text)
        return text.strip()

    @staticmethod
    def fix_unicode(text: str) -> str:
        """Fix unicode issues and normalize text."""
        return unicodedata.normalize('NFKC', text)

    @staticmethod
    def remove_special_characters(text: str) -> str:
        """Remove unwanted special characters but keep punctuation."""
        # Keep basic punctuation but remove other special characters
        return re.sub(r'[^\w\s.,!?;:()\-\'\"]+', ' ', text)

    @staticmethod
    def fix_common_issues(text: str) -> str:
        """Fix common text issues."""
        # Fix common abbreviations
        text = re.sub(r'(?<=\w)\.(?=\w)', '. ', text)
        # Fix spacing around punctuation
        text = re.sub(r'\s*([.,!?;:])\s*', r'\1 ', text)
        # Fix multiple periods
        text = re.sub(r'\.{2,}', '.', text)
        return text

    @staticmethod
    def clean_text(text: str) -> str:
        """Apply all cleaning operations."""
        text = TextCleaner.fix_unicode(text)
        text = TextCleaner.remove_special_characters(text)
        text = TextCleaner.fix_common_issues(text)
        text = TextCleaner.remove_extra_whitespace(text)
        return text

class TextChunker:
    def __init__(self, max_tokens: int = 512, overlap_tokens: int = 50):
        self.max_tokens = max_tokens
        self.overlap_tokens = overlap_tokens
        self.tokenizer = GPT2TokenizerFast.from_pretrained("gpt2")

    def count_tokens(self, text: str) -> int:
        """Count the number of tokens in a text."""
        return len(self.tokenizer(text)["input_ids"])

    def create_chunks(self, text: str) -> List[str]:
        """Create overlapping chunks of text based on semantic boundaries."""
        chunks = []
        
        # First, split into paragraphs
        paragraphs = text.split('\n\n')
        current_chunk = []
        current_tokens = 0
        last_overlap = ""
        
        for paragraph in paragraphs:
            # Clean and split paragraph into sentences
            sentences = re.split(r'(?<=[.!?])\s+', paragraph.strip())
            
            for sentence in sentences:
                sentence = sentence.strip()
                if not sentence:
                    continue
                
                sentence_tokens = self.count_tokens(sentence)
                
                # If single sentence is too long, split it
                if sentence_tokens > self.max_tokens:
                    if current_chunk:
                        chunks.append(' '.join(current_chunk))
                    words = sentence.split()
                    current_chunk = []
                    current_tokens = 0
                    temp_chunk = []
                    temp_tokens = 0
                    
                    for word in words:
                        word_tokens = self.count_tokens(word)
                        if temp_tokens + word_tokens <= self.max_tokens:
                            temp_chunk.append(word)
                            temp_tokens += word_tokens
                        else:
                            chunks.append(' '.join(temp_chunk))
                            temp_chunk = [word]
                            temp_tokens = word_tokens
                    
                    if temp_chunk:
                        current_chunk = temp_chunk
                        current_tokens = temp_tokens
                    continue
                
                # Check if adding this sentence would exceed the limit
                if current_tokens + sentence_tokens > self.max_tokens:
                    if current_chunk:
                        chunk_text = ' '.join(current_chunk)
                        chunks.append(chunk_text)
                        
                        # Create overlap with last few sentences
                        overlap_text = ' '.join(current_chunk[-2:])  # Last 2 sentences
                        if self.count_tokens(overlap_text) <= self.overlap_tokens:
                            last_overlap = overlap_text
                        else:
                            last_overlap = current_chunk[-1]  # Just the last sentence
                        
                        current_chunk = []
                        if last_overlap:
                            current_chunk.extend(last_overlap.split())
                            current_tokens = self.count_tokens(last_overlap)
                        else:
                            current_tokens = 0
                
                current_chunk.append(sentence)
                current_tokens += sentence_tokens
        
        # Add the last chunk if it exists
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks

    @staticmethod
    def read_file_content(file_path: str) -> str:
        """Read content from either .txt or .docx files."""
        if file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            # Extract text from paragraphs and tables
            text_content = []
            
            # Get text from paragraphs
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            
            # Get text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text_content.append(' | '.join(row_text))
            
            return '\n\n'.join(text_content)
        else:
            raise ValueError(f"Unsupported file format: {file_path}")

    def save_chunks(self, chunks: List[str], output_folder: str, base_name: str):
        """Save each chunk as a separate file."""
        for i, chunk in enumerate(chunks, 1):
            # Create filename with padding for proper sorting
            chunk_filename = f"{base_name}_chunk_{i:04d}.txt"
            chunk_path = os.path.join(output_folder, chunk_filename)
            
            with open(chunk_path, 'w', encoding='utf-8') as f:
                f.write(chunk.strip())

def process_files_in_folder(folder_path: str, output_dir: str, max_tokens: int = 512) -> Dict[str, List[str]]:
    """Process all text and docx files in a folder with chunking."""
    chunker = TextChunker(max_tokens=max_tokens)
    chunks_by_file = {}
    
    print(f"\nProcessing files in {folder_path}")
    print("=" * 50)

    if not os.path.exists(folder_path):
        print(f"Error: Folder not found at {folder_path}")
        return {}

    # Create chunks directory with timestamp in the specified output directory
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    chunks_folder = os.path.join(output_dir, f"chunks_{timestamp}")
    os.makedirs(chunks_folder, exist_ok=True)

    files_processed = 0
    total_chunks = 0

    for root, _, files in os.walk(folder_path):
        supported_files = [f for f in files if f.endswith(('.txt', '.docx'))]
        if not supported_files:
            print(f"No .txt or .docx files found in {folder_path}")
            continue

        print(f"\nFound {len(supported_files)} files to process")
        
        for file in supported_files:
            file_path = os.path.join(root, file)
            try:
                print(f"\nProcessing: {file}")
                
                text = TextChunker.read_file_content(file_path)
                print(f"- Read {len(text)} characters")
                
                chunks = chunker.create_chunks(text)
                print(f"- Generated {len(chunks)} initial chunks")
                
                valid_chunks = []
                for i, chunk in enumerate(chunks):
                    tokens = chunker.count_tokens(chunk)
                    if tokens <= max_tokens:
                        valid_chunks.append(chunk)
                    else:
                        print(f"  Warning: Chunk {i+1} exceeded token limit ({tokens} tokens)")
                
                # Create a subdirectory for this file's chunks
                base_name = os.path.splitext(file)[0]
                file_chunks_dir = os.path.join(chunks_folder, base_name)
                os.makedirs(file_chunks_dir, exist_ok=True)
                
                # Save individual chunk files
                chunker.save_chunks(valid_chunks, file_chunks_dir, base_name)
                
                chunks_by_file[file] = valid_chunks
                files_processed += 1
                total_chunks += len(valid_chunks)
                
                print(f"✓ Created {len(valid_chunks)} valid chunks")
                print(f"✓ Saved chunks to: {file_chunks_dir}")
                
            except Exception as e:
                print(f"Error processing {file}: {str(e)}")
                continue

    print("\nProcessing Summary:")
    print(f"Files processed: {files_processed}")
    print(f"Total chunks created: {total_chunks}")
    print(f"Chunks saved to: {chunks_folder}")
    
    return chunks_by_file

def display_chunks(chunks_by_file: Dict[str, List[str]], num_files: int = 1, num_chunks: int = 3):
    """Display processed chunks with detailed information."""
    if not chunks_by_file:
        print("No chunks to display!")
        return

    print("\nChunk Analysis:")
    print("=" * 50)
    
    for i, (filename, chunks) in enumerate(chunks_by_file.items()):
        if i >= num_files:
            break
            
        print(f"\nFile: {filename}")
        print(f"Total chunks: {len(chunks)}")
        
        chunker = TextChunker()
        for j, chunk in enumerate(chunks[:num_chunks]):
            token_count = chunker.count_tokens(chunk)
            print(f"\nChunk {j + 1}:")
            print("-" * 40)
            print(chunk)
            print("-" * 40)
            print(f"Statistics:")
            print(f"- Tokens: {token_count}")
            print(f"- Words: {len(chunk.split())}")

if __name__ == "__main__":
    # Specify your input and output directories
    input_dir = "/Users/adityabhaskara/Downloads/data"  # Replace with your input directory path
    output_dir = "/Users/adityabhaskara/Downloads/newchunks"  # Replace with your output directory path
    
    print("\nStarting chunk processing...")
    print(f"Input directory: {input_dir}")
    print(f"Output directory: {output_dir}")
    
    try:
        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
        
        # Process the files with specified output directory
        chunks_by_file = process_files_in_folder(input_dir, output_dir)
        
        if chunks_by_file:
            # Display sample chunks for verification
            print("\nDisplaying sample chunks for verification:")
            display_chunks(chunks_by_file, num_files=2, num_chunks=2)
            
            # Print summary
            total_files = len(chunks_by_file)
            total_chunks = sum(len(chunks) for chunks in chunks_by_file.values())
            print("\nProcessing Summary:")
            print(f"Total files processed: {total_files}")
            print(f"Total chunks created: {total_chunks}")
            print("\nChunk processing completed successfully!")
        else:
            print("\nNo chunks were processed. Please check if the input directory contains .txt or .docx files.")
    
    except KeyboardInterrupt:
        print("\nProcessing interrupted by user.")
    except Exception as e:
        print(f"\nAn error occurred: {str(e)}")
    finally:
        print("\nProcessing finished.")

